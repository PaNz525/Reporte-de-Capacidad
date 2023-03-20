# Import docx NOT python-docx
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_text(doc, num_vms):
    paragraph = doc.add_paragraph("Dashboard monitoreado: {{ some_text }}")
    paragraph.style = doc.styles['Heading 3']
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(5)

    paragraph = doc.add_paragraph("{{ placeholder_1 }}")
    paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(num_vms):
        paragraph = doc.add_paragraph("{{ metric_name }} – {{ resource_name }}")
        paragraph.style = doc.styles['Heading 3']
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)
    
        doc = add_table(doc)

        paragraph = doc.add_paragraph("{{ placeholder_1 }}")
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(5)
        paragraph.paragraph_format.space_after = Pt(5)

    return doc

def add_table(doc):
    # Table headers
    table_headers = ["Min", "Max", "Mean"]

    # Table data in a form of list
    peformance_values = [
        ["{{ min_value }}", "{{ max_value }}", "{{ mean_value }}"]
    ]
    
    # Creating a table object
    table = doc.add_table(rows=1, cols=3)
    
    for i in range(3):
        table.rows[0].cells[i].text = table_headers[i]
    
    # Adding data from the list to the table
    for min, max, mean in peformance_values:
        cells = table.add_row().cells
        cells[0].text = min
        cells[1].text = max
        cells[2].text = mean
    
    # Adding style to a table
    table.style = 'Medium Grid 1 Accent 1'

    return doc

def main():
    # Create an instance of a word document
    doc = docx.Document()

    add_text(doc, 5)

    # Now save the document to a location
    doc.save('python_template.docx')

if __name__ == "__main__":
    main()